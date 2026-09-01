import io
import json
import uuid

import pytest
from docx import Document
from fakes import VALID_PROFILE, ProviderError, install_acompletion, llm_response
from httpx import ASGITransport, AsyncClient
from sqlalchemy import func, select

from app.core.config import get_settings
from app.core.db import session_factory
from app.main import app
from app.models import Candidate, Resume

pytestmark = pytest.mark.usefixtures("clean_tables")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXPECTED_PARSE_VERSION = "gemini/gemini-2.5-flash+profile_prompt_v5"


@pytest.fixture(autouse=True)
def gemini_key(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("GEMINI_API_KEY", "test-key")
    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture
async def client() -> AsyncClient:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://testserver") as async_client:
        yield async_client


def docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Data Analyst with 8 years of experience")
    document.save(buffer)
    return buffer.getvalue()


async def upload_resume(client: AsyncClient) -> dict:
    response = await client.post(
        "/api/resumes", files={"file": ("resume.docx", io.BytesIO(docx_bytes()), DOCX_MIME)}
    )
    assert response.status_code == 201
    return response.json()


def disable_sources(monkeypatch: pytest.MonkeyPatch) -> None:
    settings = get_settings()
    monkeypatch.setattr(settings, "adzuna_app_id", None)
    monkeypatch.setattr(settings, "adzuna_app_key", None)
    monkeypatch.setattr(settings, "apify_token", None)


async def test_extract_returns_draft_and_persists_parse_artifact(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    disable_sources(monkeypatch)
    uploaded = await upload_resume(client)
    calls = install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps(VALID_PROFILE)))

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 200
    body = response.json()
    assert body["resume_id"] == uploaded["resume_id"]
    assert body["candidate_id"] == uploaded["candidate_id"]
    assert body["draft_profile"]["contact"]["full_name"] == "Jane Doe"
    assert body["draft_profile"]["experience"][0]["is_current"] is True
    assert body["draft_profile"]["contact"]["links"][0]["label"] == "LinkedIn"
    assert body["draft_profile"]["contact"]["links"][1]["label"] == "GitHub"
    assert body["draft_profile"]["contact"]["links"][2]["label"] == "Website"
    assert body["draft_profile"]["projects"][0]["name"] == "OpenPipeline"
    assert body["draft_profile"]["awards"][0]["title"] == "Winner, HackX 2023"
    assert body["draft_profile"]["extra_sections"][0]["title"] == "Publications"
    assert body["parse_version"] == EXPECTED_PARSE_VERSION
    assert len(calls) == 1
    assert "Senior Data Analyst" in calls[0]["messages"][1]["content"]

    async with session_factory() as session:
        resume = await session.get(Resume, uuid.UUID(uploaded["resume_id"]))
        candidate_count = (
            await session.execute(select(func.count()).select_from(Candidate))
        ).scalar_one()
    assert resume is not None
    assert resume.draft_profile is not None
    assert resume.draft_profile["contact"]["full_name"] == "Jane Doe"
    assert resume.draft_profile["extra_sections"][0]["entries"] == [
        "Doe J. Efficient Pipelines, 2022"
    ]
    assert resume.parsed_at is not None
    assert resume.parse_version == EXPECTED_PARSE_VERSION
    assert candidate_count == 1


async def test_reextract_overwrites_draft(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    disable_sources(monkeypatch)
    uploaded = await upload_resume(client)
    install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps(VALID_PROFILE)))
    first = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")
    assert first.status_code == 200

    updated = {**VALID_PROFILE, "headline": "Principal Analyst"}
    install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps(updated)))
    second = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert second.status_code == 200
    assert second.json()["draft_profile"]["headline"] == "Principal Analyst"


async def test_unknown_resume_returns_404(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps(VALID_PROFILE)))

    response = await client.post(f"/api/resumes/{uuid.uuid4()}/extract")

    assert response.status_code == 404


async def test_resume_without_text_returns_409(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async with session_factory() as session:
        candidate = Candidate()
        session.add(candidate)
        await session.flush()
        resume = Resume(
            candidate_id=candidate.id,
            file_path="unused.pdf",
            original_filename="x.pdf",
            content_type="application/pdf",
            size_bytes=1,
            extracted_text=None,
        )
        session.add(resume)
        await session.flush()
        resume_id = resume.id
        await session.commit()
    install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps(VALID_PROFILE)))

    response = await client.post(f"/api/resumes/{resume_id}/extract")

    assert response.status_code == 409


async def test_repair_exhausted_returns_502_and_leaves_row_untouched(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    uploaded = await upload_resume(client)
    install_acompletion(monkeypatch, lambda **kw: llm_response(json.dumps({"skills": ["SQL"]})))

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 502
    assert response.json()["detail"].startswith("structured output failed validation after repair")
    async with session_factory() as session:
        resume = await session.get(Resume, uuid.UUID(uploaded["resume_id"]))
    assert resume is not None
    assert resume.draft_profile is None
    assert resume.parse_version == "text_v1"
    assert resume.parsed_at is not None


async def test_unconfigured_llm_returns_503(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    def not_configured() -> bool:
        return False

    monkeypatch.setattr("app.services.profile_extraction.is_llm_configured", not_configured)
    uploaded = await upload_resume(client)

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 503


async def test_transport_failure_detail_carries_cause_hint(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def no_delay(_: float) -> None:
        return None

    monkeypatch.setattr("app.adapters.llm.asyncio.sleep", no_delay)
    uploaded = await upload_resume(client)
    responses = iter([ProviderError(429), ProviderError(429)])
    install_acompletion(monkeypatch, lambda **kw: next(responses))

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 502
    assert "rate limited by the provider" in response.json()["detail"]


def routed_handler(queries_payload: object):
    def handler(**kw: object) -> object:
        if kw.get("temperature") == 0.8:
            return llm_response(json.dumps(queries_payload))
        return llm_response(json.dumps(VALID_PROFILE))

    return handler


QUERIES_PAYLOAD = {
    "queries": {
        "adzuna": {"title": "Senior Data Analyst", "skills": ["SQL", "Python"], "exclude": []}
    }
}


async def test_extract_generates_search_queries(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(get_settings(), "adzuna_app_id", "id")
    monkeypatch.setattr(get_settings(), "adzuna_app_key", "key")
    calls = install_acompletion(monkeypatch, routed_handler(QUERIES_PAYLOAD))
    uploaded = await upload_resume(client)

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 200
    body = response.json()
    assert body["search_queries"]["queries"]["adzuna"]["title"] == "Senior Data Analyst"
    assert body["search_queries"]["prompt_version"] == "search_query_v1"
    temperatures = [call["temperature"] for call in calls]
    assert 0.2 in temperatures and 0.8 in temperatures
    query_prompt = next(call for call in calls if call["temperature"] == 0.8)
    assert "resume text" not in query_prompt["messages"][1]["content"].lower()
    async with session_factory() as session:
        resume = await session.get(Resume, uuid.UUID(uploaded["resume_id"]))
    assert resume is not None
    assert resume.search_queries is not None
    assert resume.search_queries["queries"]["adzuna"]["title"] == "Senior Data Analyst"


async def test_extract_survives_query_generation_failure(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(get_settings(), "adzuna_app_id", "id")
    monkeypatch.setattr(get_settings(), "adzuna_app_key", "key")

    async def no_delay(_: float) -> None:
        return None

    def handler(**kw: object) -> object:
        if kw.get("temperature") == 0.8:
            return ProviderError(429)
        return llm_response(json.dumps(VALID_PROFILE))

    monkeypatch.setattr("app.adapters.llm.asyncio.sleep", no_delay)
    install_acompletion(monkeypatch, handler)
    uploaded = await upload_resume(client)

    response = await client.post(f"/api/resumes/{uploaded['resume_id']}/extract")

    assert response.status_code == 200
    assert response.json()["search_queries"] is None
    async with session_factory() as session:
        resume = await session.get(Resume, uuid.UUID(uploaded["resume_id"]))
    assert resume is not None
    assert resume.draft_profile is not None
    assert resume.search_queries is None
