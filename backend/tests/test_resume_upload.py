import io
import uuid
from pathlib import Path

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient
from sqlalchemy import select

from app.core.db import session_factory
from app.main import app
from app.models import Resume

pytestmark = pytest.mark.usefixtures("clean_tables")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture
async def client() -> AsyncClient:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://testserver") as async_client:
        yield async_client


def docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Data Analyst with 8 years of experience")
    document.save(buffer)
    return buffer.getvalue()


async def test_upload_docx_persists_metadata_and_returns_text(client: AsyncClient) -> None:
    response = await client.post(
        "/api/resume",
        files={"file": ("jane_resume.docx", io.BytesIO(docx_bytes()), DOCX_MIME)},
    )

    assert response.status_code == 201
    body = response.json()
    uuid.UUID(body["resume_id"])
    uuid.UUID(body["candidate_id"])
    assert body["original_filename"] == "jane_resume.docx"
    assert "Senior Data Analyst" in body["extracted_text"]
    assert body["page_count"] is None
    assert body["parse_version"] == "text_v1"
    assert body["size_bytes"] > 0

    async with session_factory() as session:
        row = (await session.execute(select(Resume))).scalars().one()
    assert row.original_filename == "jane_resume.docx"
    saved_path = Path(row.file_path)
    assert saved_path.exists()
    assert saved_path.suffix == ".docx"
    uuid.UUID(saved_path.stem)


async def test_upload_pdf_returns_page_count(client: AsyncClient) -> None:
    data = (FIXTURES / "two_page_resume.pdf").read_bytes()
    response = await client.post(
        "/api/resume", files={"file": ("two_page.pdf", io.BytesIO(data), "application/pdf")}
    )

    assert response.status_code == 201
    body = response.json()
    assert body["page_count"] == 2
    assert "Jane Doe" in body["extracted_text"]


async def test_second_upload_reuses_single_candidate(client: AsyncClient) -> None:
    docx_response = await client.post(
        "/api/resume", files={"file": ("a.docx", io.BytesIO(docx_bytes()), DOCX_MIME)}
    )
    pdf_data = (FIXTURES / "single_column_resume.pdf").read_bytes()
    pdf_response = await client.post(
        "/api/resume", files={"file": ("b.pdf", io.BytesIO(pdf_data), "application/pdf")}
    )

    assert docx_response.status_code == 201
    assert pdf_response.status_code == 201
    assert docx_response.json()["candidate_id"] == pdf_response.json()["candidate_id"]


async def test_oversized_upload_rejected_with_413(client: AsyncClient) -> None:
    oversized = b"%PDF-1.4\n" + b"0" * (10 * 1024 * 1024 + 1)
    response = await client.post(
        "/api/resume", files={"file": ("big.pdf", io.BytesIO(oversized), "application/pdf")}
    )

    assert response.status_code == 413
    assert "detail" in response.json()


async def test_unsupported_type_rejected_with_415(client: AsyncClient) -> None:
    response = await client.post(
        "/api/resume",
        files={"file": ("notes.txt", io.BytesIO(b"plain text"), "text/plain")},
    )

    assert response.status_code == 415
    assert "PDF and DOCX" in response.json()["detail"]


async def test_extension_magic_mismatch_rejected_with_415(client: AsyncClient) -> None:
    response = await client.post(
        "/api/resume",
        files={"file": ("fake.pdf", io.BytesIO(b"PK\x03\x04 not a pdf"), "application/pdf")},
    )

    assert response.status_code == 415


async def test_pdf_without_text_rejected_with_422(client: AsyncClient) -> None:
    data = (FIXTURES / "empty_resume.pdf").read_bytes()
    response = await client.post(
        "/api/resume", files={"file": ("scanned.pdf", io.BytesIO(data), "application/pdf")}
    )

    assert response.status_code == 422
    assert "no readable text" in response.json()["detail"]
