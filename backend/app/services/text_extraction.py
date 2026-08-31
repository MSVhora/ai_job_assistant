import io
import logging
from pathlib import Path
from typing import Literal

import pdfplumber
from docx import Document

from app.core.errors import TextExtractionError, UnsupportedFileTypeError

logger = logging.getLogger(__name__)

_PDF_MAGIC = b"%PDF-"
_ZIP_MAGIC = b"PK\x03\x04"

SupportedKind = Literal["pdf", "docx"]


def sniff_file_type(filename: str, head: bytes) -> SupportedKind:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf" and head.startswith(_PDF_MAGIC):
        return "pdf"
    if extension == ".docx" and head.startswith(_ZIP_MAGIC):
        return "docx"
    raise UnsupportedFileTypeError()


def extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            page_texts = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:
        logger.warning("pdf text extraction failed", exc_info=True)
        raise TextExtractionError("file could not be parsed as a PDF") from exc
    return "\n\n".join(t for t in page_texts if t.strip()), page_count


def extract_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        logger.warning("docx text extraction failed", exc_info=True)
        raise TextExtractionError("file could not be parsed as a DOCX") from exc
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)
