import io
from pathlib import Path

import pytest
from docx import Document

from app.core.errors import TextExtractionError, UnsupportedFileTypeError
from app.services.text_extraction import extract_docx, extract_pdf, sniff_file_type

FIXTURES = Path(__file__).parent / "fixtures"


def build_docx(paragraphs: list[str], table_cells: list[list[str]] | None = None) -> bytes:
    buffer = io.BytesIO()
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_cells:
        table = document.add_table(rows=len(table_cells), cols=len(table_cells[0]))
        for r, row in enumerate(table_cells):
            for c, cell_text in enumerate(row):
                table.rows[r].cells[c].text = cell_text
    document.save(buffer)
    return buffer.getvalue()


def test_sniff_accepts_matching_extension_and_magic() -> None:
    assert sniff_file_type("resume.pdf", b"%PDF-1.4\n") == "pdf"
    assert sniff_file_type("resume.docx", b"PK\x03\x04....") == "docx"


def test_sniff_rejects_wrong_extension() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        sniff_file_type("notes.txt", b"hello")
    with pytest.raises(UnsupportedFileTypeError):
        sniff_file_type("legacy.doc", b"PK\x03\x04")


def test_sniff_rejects_content_that_does_not_match_extension() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        sniff_file_type("fake.pdf", b"PK\x03\x04not a pdf")
    with pytest.raises(UnsupportedFileTypeError):
        sniff_file_type("empty.pdf", b"")


def test_extract_pdf_single_column_fixture() -> None:
    text, page_count = extract_pdf((FIXTURES / "single_column_resume.pdf").read_bytes())
    assert page_count == 1
    assert "Jane Doe" in text
    assert "Senior Data Analyst" in text
    assert "Skills: SQL, Python, Tableau, dbt" in text


def test_extract_pdf_multi_page_counts_pages() -> None:
    text, page_count = extract_pdf((FIXTURES / "two_page_resume.pdf").read_bytes())
    assert page_count == 2
    assert "Jane Doe" in text
    assert "Meridian Retail" in text


def test_extract_pdf_without_text_returns_empty() -> None:
    text, page_count = extract_pdf((FIXTURES / "empty_resume.pdf").read_bytes())
    assert text == ""
    assert page_count == 1


def test_extract_pdf_garbage_raises_domain_error() -> None:
    with pytest.raises(TextExtractionError):
        extract_pdf(b"%PDF-1.4 this is not a real pdf body")


def test_extract_docx_paragraphs_and_tables() -> None:
    data = build_docx(
        ["Jane Doe", "Senior Data Analyst"],
        table_cells=[["Company", "Years"], ["Meridian Retail", "5"]],
    )
    text = extract_docx(data)
    assert "Jane Doe" in text
    assert "Meridian Retail" in text


def test_extract_docx_garbage_raises_domain_error() -> None:
    with pytest.raises(TextExtractionError):
        extract_docx(b"PK\x03\x04 definitely not a docx archive")
